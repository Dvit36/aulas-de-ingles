"""Safe validation, text extraction and storage for non-image evidence files."""

from __future__ import annotations

from dataclasses import dataclass
from hashlib import sha256
from io import BytesIO
from math import ceil, isfinite
import os
from pathlib import Path
from uuid import uuid4
from zipfile import BadZipFile, ZipFile


SUPPORTED_EXTENSIONS = frozenset({".png", ".jpg", ".jpeg", ".webp", ".pdf", ".docx", ".txt"})
DOCUMENT_EXTENSIONS = frozenset({".pdf", ".docx", ".txt"})
MAX_DOCX_ENTRIES = 1_000
MAX_DOCX_COMPRESSION_RATIO = 200.0
PDF_RENDER_SCALE = 1.6


class DocumentValidationError(ValueError):
    def __init__(self, message: str, *, code: str = "invalid_document") -> None:
        super().__init__(message)
        self.message = message
        self.code = code


@dataclass(frozen=True)
class ProcessedDocument:
    original_bytes: bytes
    file_kind: str
    mime_type: str
    extension: str
    byte_size: int
    sha256: str
    extracted_text: str
    page_count: int | None = None


def _extract_pdf(
    payload: bytes,
    *,
    max_pages: int,
    max_render_pixels: int,
    ocr_engine=None,
) -> tuple[str, int]:
    from pypdf import PdfReader

    try:
        reader = PdfReader(BytesIO(payload), strict=True)
    except Exception as error:
        raise DocumentValidationError("PDF inválido ou corrompido", code="invalid_pdf") from error
    if reader.is_encrypted:
        raise DocumentValidationError("PDF protegido por senha não é aceito", code="encrypted_pdf")
    page_count = len(reader.pages)
    if page_count < 1 or page_count > max_pages:
        raise DocumentValidationError(
            f"O PDF deve ter entre 1 e {max_pages} páginas",
            code="pdf_page_limit",
        )
    estimated_render_pixels = 0
    for page in reader.pages:
        try:
            width_points = float(page.mediabox.width)
            height_points = float(page.mediabox.height)
        except (TypeError, ValueError, OverflowError) as error:
            raise DocumentValidationError(
                "PDF possui dimensões de página inválidas",
                code="pdf_invalid_dimensions",
            ) from error
        if (
            not isfinite(width_points)
            or not isfinite(height_points)
            or width_points <= 0
            or height_points <= 0
        ):
            raise DocumentValidationError(
                "PDF possui dimensões de página inválidas",
                code="pdf_invalid_dimensions",
            )
        estimated_render_pixels += ceil(width_points * PDF_RENDER_SCALE) * ceil(
            height_points * PDF_RENDER_SCALE
        )
        if estimated_render_pixels > max_render_pixels:
            raise DocumentValidationError(
                "PDF excede o limite seguro de renderização",
                code="pdf_render_limit",
            )
    text = "\n\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
    if text or ocr_engine is None:
        return text, page_count

    # Scanned PDF fallback. The renderer and OCR are loaded only for this case.
    import pypdfium2 as pdfium

    from .ocr import extract_text

    document = pdfium.PdfDocument(payload)
    chunks: list[str] = []
    try:
        for page_index in range(page_count):
            page = document[page_index]
            try:
                bitmap = page.render(scale=PDF_RENDER_SCALE)
                try:
                    image = bitmap.to_pil()
                    output = BytesIO()
                    image.save(output, format="PNG")
                    result = extract_text(output.getvalue(), engine=ocr_engine)
                    if result.text.strip():
                        chunks.append(result.text.strip())
                finally:
                    bitmap.close()
            finally:
                page.close()
    finally:
        document.close()
    return "\n\n".join(chunks), page_count


def _extract_docx(payload: bytes, *, max_expanded_bytes: int) -> str:
    try:
        with ZipFile(BytesIO(payload)) as archive:
            members = archive.infolist()
            if len(members) > MAX_DOCX_ENTRIES:
                raise DocumentValidationError(
                    "DOCX contém arquivos internos demais",
                    code="docx_entry_limit",
                )
            expanded_bytes = 0
            for member in members:
                if member.flag_bits & 0x1:
                    raise DocumentValidationError(
                        "DOCX criptografado não é aceito",
                        code="encrypted_docx",
                    )
                expanded_bytes += int(member.file_size)
                if expanded_bytes > max_expanded_bytes:
                    raise DocumentValidationError(
                        "DOCX excede o limite seguro após descompactação",
                        code="docx_expanded_limit",
                    )
                if member.file_size and (
                    member.file_size / max(1, member.compress_size)
                    > MAX_DOCX_COMPRESSION_RATIO
                ):
                    raise DocumentValidationError(
                        "DOCX possui taxa de compressão insegura",
                        code="docx_compression_ratio",
                    )
            names = {member.filename for member in members}
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise DocumentValidationError("O arquivo não é um DOCX válido", code="invalid_docx")
            lowered = {name.casefold() for name in names}
            if any("vbaproject.bin" in name or "/embeddings/" in name for name in lowered):
                raise DocumentValidationError(
                    "DOCX com macro ou objeto incorporado não é aceito",
                    code="unsafe_docx",
                )
    except BadZipFile as error:
        raise DocumentValidationError("DOCX inválido ou corrompido", code="invalid_docx") from error

    from docx import Document

    try:
        document = Document(BytesIO(payload))
    except Exception as error:
        raise DocumentValidationError("Não foi possível ler o DOCX", code="invalid_docx") from error
    chunks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            line = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if line:
                chunks.append(line)
    return "\n".join(chunks)


def _extract_txt(payload: bytes) -> str:
    if b"\x00" in payload:
        raise DocumentValidationError("TXT binário não é aceito", code="binary_txt")
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentValidationError("Não foi possível identificar o encoding do TXT", code="txt_encoding")


def process_document_bytes(
    payload: bytes,
    filename: str,
    *,
    max_bytes: int,
    max_pdf_pages: int,
    max_document_expanded_bytes: int = 50 * 1024 * 1024,
    max_pdf_render_pixels: int = 24_000_000,
    ocr_engine=None,
) -> ProcessedDocument:
    if not payload:
        raise DocumentValidationError("Arquivo vazio", code="empty_file")
    if len(payload) > max_bytes:
        raise DocumentValidationError("Arquivo excede o limite configurado", code="file_too_large")
    extension = Path(filename.replace("\x00", "")).suffix.casefold()
    if extension == ".doc":
        raise DocumentValidationError(
            "Arquivos .doc antigos não são aceitos; converta para .docx",
            code="legacy_doc",
        )
    if extension not in DOCUMENT_EXTENSIONS:
        raise DocumentValidationError("Tipo de documento não suportado", code="unsupported_type")

    if payload.startswith(b"%PDF-"):
        detected_extension, kind, mime = ".pdf", "pdf", "application/pdf"
        text, page_count = _extract_pdf(
            payload,
            max_pages=max_pdf_pages,
            max_render_pixels=max_pdf_render_pixels,
            ocr_engine=ocr_engine,
        )
    elif payload.startswith(b"PK\x03\x04"):
        detected_extension, kind, mime = (
            ".docx",
            "docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        text, page_count = _extract_docx(
            payload, max_expanded_bytes=max_document_expanded_bytes
        ), None
    else:
        detected_extension, kind, mime = ".txt", "txt", "text/plain"
        text, page_count = _extract_txt(payload), None
    if extension != detected_extension:
        raise DocumentValidationError(
            "A extensão não corresponde ao conteúdo real do arquivo",
            code="extension_mismatch",
        )
    return ProcessedDocument(
        original_bytes=payload,
        file_kind=kind,
        mime_type=mime,
        extension=extension,
        byte_size=len(payload),
        sha256=sha256(payload).hexdigest(),
        extracted_text=text.strip(),
        page_count=page_count,
    )


def persist_document(document: ProcessedDocument, upload_directory: Path) -> tuple[str, Path]:
    upload_directory.mkdir(parents=True, exist_ok=True, mode=0o700)
    try:
        upload_directory.chmod(0o700)
    except OSError:
        pass
    storage_key = f"{uuid4().hex}{document.extension}"
    path = upload_directory / storage_key
    descriptor = os.open(path, os.O_WRONLY | os.O_CREAT | os.O_EXCL, 0o600)
    try:
        with os.fdopen(descriptor, "wb") as handle:
            handle.write(document.original_bytes)
            handle.flush()
            os.fsync(handle.fileno())
    except Exception:
        path.unlink(missing_ok=True)
        raise
    return storage_key, path


__all__ = [
    "DOCUMENT_EXTENSIONS",
    "SUPPORTED_EXTENSIONS",
    "DocumentValidationError",
    "ProcessedDocument",
    "persist_document",
    "process_document_bytes",
]
